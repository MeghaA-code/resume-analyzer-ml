"""
text_extraction.py
-------------------
Extracts raw text from an uploaded resume file (PDF or DOCX).

Two extraction backends are used for PDFs:
  1. pdfplumber (better layout / table handling)
  2. pypdf as a fallback if pdfplumber fails or returns empty text

DOCX files are parsed with python-docx, pulling text from paragraphs
and table cells (some resumes use tables for layout).
"""

import os
import docx
import pdfplumber
from pypdf import PdfReader


class TextExtractionError(Exception):
    """Raised when a resume file cannot be parsed into text."""


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file, trying pdfplumber first, then pypdf."""
    text_chunks = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_chunks.append(page_text)
        text = "\n".join(text_chunks).strip()
        if text:
            return text
    except Exception:
        pass  # fall back to pypdf below

    try:
        reader = PdfReader(file_path)
        text_chunks = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(text_chunks).strip()
        if text:
            return text
    except Exception as exc:
        raise TextExtractionError(f"Failed to read PDF: {exc}") from exc

    raise TextExtractionError(
        "No extractable text found in PDF (it may be a scanned image)."
    )


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file, including paragraphs and table cells."""
    try:
        document = docx.Document(file_path)
    except Exception as exc:
        raise TextExtractionError(f"Failed to read DOCX: {exc}") from exc

    text_chunks = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_chunks.append(cell.text)

    text = "\n".join(text_chunks).strip()
    if not text:
        raise TextExtractionError("No extractable text found in DOCX file.")
    return text


def extract_text(file_path: str) -> str:
    """Dispatch to the correct extractor based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)

    raise TextExtractionError(f"Unsupported file type: {ext}")
